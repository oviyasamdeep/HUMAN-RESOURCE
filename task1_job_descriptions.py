from docx import Document

def create_job_description(role, responsibilities, qualifications, skills, filename):
    doc = Document()
    doc.add_heading(f'Job Description: {role}', level=1)

    doc.add_heading('Responsibilities', level=2)
    for item in responsibilities:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Qualifications', level=2)
    for item in qualifications:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Required Skills', level=2)
    for item in skills:
        doc.add_paragraph(item, style='List Bullet')

    doc.save(filename)

# Software Developer
create_job_description(
    "Software Developer",
    [
        "Develop and maintain web applications",
        "Write clean and efficient code",
        "Collaborate with cross-functional teams",
        "Debug and test software"
    ],
    [
        "Bachelor’s degree in Computer Science",
        "0-2 years experience (Freshers can apply)"
    ],
    [
        "Python/Java knowledge",
        "Problem-solving ability",
        "Database knowledge",
        "Teamwork skills"
    ],
    "Software_Developer_JD.docx"
)

# HR Executive
create_job_description(
    "HR Executive",
    [
        "Manage recruitment process",
        "Maintain employee records",
        "Handle payroll coordination",
        "Organize training programs"
    ],
    [
        "Bachelor’s degree in HR/Business Administration",
        "Strong communication skills"
    ],
    [
        "MS Office proficiency",
        "Interpersonal skills",
        "Time management",
        "Organizational skills"
    ],
    "HR_Executive_JD.docx"
)

print("Task 1 Completed")
