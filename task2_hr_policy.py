from docx import Document

doc = Document()
doc.add_heading('Startup HR Policy Document', level=1)

doc.add_heading('1. Leave Policy', level=2)
doc.add_paragraph("""
Employees are entitled to:
- 12 Casual Leaves per year
- 10 Sick Leaves per year
- 15 Earned Leaves per year
Leave must be applied through HR portal.
""")

doc.add_heading('2. Code of Conduct', level=2)
doc.add_paragraph("""
Employees must:
- Maintain professional behavior
- Avoid harassment and discrimination
- Protect company confidential information
- Follow ethical practices
""")

doc.add_heading('3. Work From Home Guidelines', level=2)
doc.add_paragraph("""
- Employees must be available during work hours.
- Daily progress updates required.
- Maintain data security.
- Follow productivity standards.
""")

doc.save("HR_Policy_Document.docx")

print("Task 2 Completed")
