from docx import Document

doc = Document()
doc.add_heading('HR Trends Report: Remote Work & HR Challenges', level=1)

content = """
Remote work has transformed the modern workplace significantly...
(Write detailed 1000+ words content here covering:)
- Growth of remote work
- Benefits and challenges
- Digital HR transformation
- Employee engagement issues
- Future of hybrid model
- HR technology adaptation
"""

doc.add_paragraph(content)

doc.save("HR_Trends_Report.docx")

print("Task 4 Completed")
